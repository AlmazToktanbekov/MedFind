"""Создаёт docs/vkr/_шаблон_гост.docx — reference-doc для pandoc.
Оформление по методичке ВКР КГТУ им. И. Раззакова:
A4, поля 30/10/20/20 мм, Times New Roman 14, интервал 1.5,
выравнивание по ширине, абзацный отступ 12.5 мм,
заголовки глав прописными по центру, нумерация страниц снизу по центру.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent / "_шаблон_гост.docx"
FONT = "Times New Roman"

doc = Document()

# ── Формат страницы: A4 + поля ───────────────────────────────────────────────
for s in doc.sections:
    s.page_width = Mm(210)
    s.page_height = Mm(297)
    s.left_margin = Mm(30)
    s.right_margin = Mm(10)
    s.top_margin = Mm(20)
    s.bottom_margin = Mm(20)


def set_font(style, name=FONT):
    """Прописывает шрифт во все слоты rFonts (для корректной кириллицы)."""
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for tag in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(tag), name)


# ── Стиль обычного текста ────────────────────────────────────────────────────
normal = doc.styles["Normal"]
set_font(normal)
normal.font.size = Pt(14)
pf = normal.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.first_line_indent = Mm(12.5)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ── Стили заголовков ─────────────────────────────────────────────────────────
# уровень: (выравнивание, отступ первой строки мм, разрыв страницы до)
heading_cfg = {
    1: (WD_ALIGN_PARAGRAPH.CENTER, 0.0, True),
    2: (WD_ALIGN_PARAGRAPH.LEFT, 12.5, False),
    3: (WD_ALIGN_PARAGRAPH.LEFT, 12.5, False),
}
for lvl, (align, indent, page_break) in heading_cfg.items():
    st = doc.styles[f"Heading {lvl}"]
    set_font(st)
    st.font.size = Pt(14)
    st.font.bold = True
    st.font.italic = False
    st.font.color.rgb = RGBColor(0, 0, 0)
    hp = st.paragraph_format
    hp.alignment = align
    hp.first_line_indent = Mm(indent)
    hp.line_spacing = 1.5
    hp.space_before = Pt(12)
    hp.space_after = Pt(12)
    hp.keep_with_next = True
    hp.page_break_before = page_break

# ── Нумерация страниц: снизу по центру ───────────────────────────────────────
for s in doc.sections:
    p = s.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._element.append(begin)
    run._element.append(instr)
    run._element.append(end)
    run.font.name = FONT
    run.font.size = Pt(14)

doc.save(OUT)
print(f"Шаблон сохранён: {OUT}")
