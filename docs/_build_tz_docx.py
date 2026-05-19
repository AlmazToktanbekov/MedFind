"""Сборка docs/TZ_MedFind_v2_1.docx — детальное техническое задание.
Оформление по методичке ВКР КГТУ: Times New Roman 14, интервал 1.5,
поля 30/10/20/20 мм, нумерация страниц снизу по центру.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent / "TZ_MedFind_v2_1.docx"

doc = Document()

# ── Page setup ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.left_margin = Mm(30)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)


def set_font(run, *, bold=False, italic=False, size=14, mono=False):
    name = "Consolas" if mono else "Times New Roman"
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for tag in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(tag), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_numbers(doc):
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for typ in ("begin",):
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), typ)
        run._element.append(el)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._element.append(instr)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._element.append(end)
    set_font(run, size=12)


def p(text="", *, bold=False, italic=False, size=14, align="justify",
      indent_first=12.5, space_before=0, space_after=0, mono=False):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if indent_first:
        para.paragraph_format.first_line_indent = Mm(indent_first)
    para.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    if text:
        run = para.add_run(text)
        set_font(run, bold=bold, italic=italic, size=size, mono=mono)
    return para


def h1(text):
    doc.add_page_break()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text.upper())
    set_font(run, bold=True, size=16)


def h2(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    set_font(run, bold=True, size=14)


def h3(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    set_font(run, bold=True, italic=True, size=14)


def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    set_font(run, size=14)


_num_counter = [0]


def reset_numbering():
    _num_counter[0] = 0


def numbered(text):
    _num_counter[0] += 1
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.left_indent = Mm(12.5)
    run = para.add_run(f"{_num_counter[0]}. {text}")
    set_font(run, size=14)


def code_block(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Mm(10)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    set_font(run, size=11, mono=True)


def table(headers, rows, widths_mm=None, font_size=11):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.autofit = False
    if widths_mm:
        for i, w in enumerate(widths_mm):
            for cell in tbl.columns[i].cells:
                cell.width = Mm(w)
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        set_font(run, bold=True, size=font_size)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_i, row in enumerate(rows, 1):
        cells = tbl.rows[r_i].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            para = cells[c_i].paragraphs[0]
            para.paragraph_format.line_spacing = 1.15
            run = para.add_run(str(val))
            set_font(run, size=font_size)
    return tbl


# ═══════════════════════════════════════════════════════════════════════════
# ТИТУЛЬНЫЙ ЛИСТ
# ═══════════════════════════════════════════════════════════════════════════
p("МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ КЫРГЫЗСКОЙ РЕСПУБЛИКИ",
  bold=True, align="center", indent_first=0, space_before=24)
p("КЫРГЫЗСКИЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ им. И. РАЗЗАКОВА",
  bold=True, align="center", indent_first=0, space_after=18)
p("Факультет информационных технологий", align="center", indent_first=0)
p("Кафедра информатики и вычислительной техники",
  align="center", indent_first=0, space_after=60)

p("«УТВЕРЖДАЮ»", align="right", indent_first=0, space_before=12)
p("Заведующий кафедрой", align="right", indent_first=0)
p("__________________________", align="right", indent_first=0)
p("«___» _____________ 2026 г.", align="right", indent_first=0, space_after=48)

p("ТЕХНИЧЕСКОЕ ЗАДАНИЕ", bold=True, size=20,
  align="center", indent_first=0, space_before=24, space_after=12)
p("на выпускную квалификационную работу", italic=True,
  align="center", indent_first=0, space_after=36)
p("Тема: «Разработка мобильного приложения «MedFind» для поиска врачей, "
  "клиник и аптек в Кыргызской Республике»",
  bold=True, align="center", indent_first=0, space_after=36)

p("Студент: Токтанбеков Алмаз Абу-Насирович", indent_first=0, space_after=6)
p("Группа: ИВТ-1-22", indent_first=0, space_after=6)
p("Направление: 710100 — Информатика и вычислительная техника (бакалавриат)",
  indent_first=0, space_after=6)
p("Руководитель: ___________________________________________________",
  indent_first=0, space_after=24)
p("Бишкек — 2026", align="center", indent_first=0, space_before=48)


# ═══════════════════════════════════════════════════════════════════════════
# СОДЕРЖАНИЕ
# ═══════════════════════════════════════════════════════════════════════════
h1("Содержание")
toc = [
    "1. Введение",
    "2. Цели и задачи проекта",
    "3. Целевая аудитория и роли",
    "4. Анализ конкурентов",
    "5. Архитектура системы",
    "6. Используемые технологии",
    "7. Структура серверной части (Backend)",
    "8. Структура мобильного клиента (Frontend / Flutter)",
    "9. База данных",
    "10. REST API — детальный перечень эндпоинтов",
    "11. Функциональные требования по ролям",
    "12. Нефункциональные требования",
    "13. Дизайн-система",
    "14. Безопасность",
    "15. Тестирование",
    "16. Развёртывание и сопровождение",
    "17. Конфигурация (.env)",
    "18. Этапы разработки и статус",
    "19. Что предстоит сделать",
    "20. Заключение",
    "21. Список использованной литературы",
]
for item in toc:
    p(item, indent_first=0)


# ═══════════════════════════════════════════════════════════════════════════
# 1. ВВЕДЕНИЕ
# ═══════════════════════════════════════════════════════════════════════════
h1("1. Введение")

h2("1.1. Краткое описание")
p("MedFind — мобильное приложение для Кыргызстана, объединяющее в одной "
  "платформе врачей, медицинские клиники и аптеки. Пациенты получают единый "
  "удобный канал для поиска медицинских услуг, а медицинские учреждения — "
  "инструмент для привлечения новых клиентов и управления своим профилем. "
  "Пациенты, врачи, клиники и аптеки пользуются платформой бесплатно.")

h2("1.2. Актуальность")
p("На текущий момент в Кыргызской Республике отсутствует специализированный "
  "мобильный сервис-агрегатор медицинских услуг. Информация о врачах и "
  "клиниках распределена между сторонними справочниками (2GIS), социальными "
  "сетями (Instagram, WhatsApp-каналы), сайтами клиник и сарафанным радио. "
  "Это создаёт следующие проблемы:")
bullet("Пациент тратит значительное время на поиск нужного специалиста и не "
       "может предварительно сравнить врачей по объективным критериям.")
bullet("Отсутствует институт верификации врачей — невозможно убедиться, что "
       "специалист действительно работает в указанной клинике.")
bullet("Медицинские учреждения не имеют доступного цифрового канала "
       "привлечения клиентов, кроме рекламы в социальных сетях.")
bullet("Аптеки не имеют единой витрины — пациент не может удобно найти "
       "ближайший филиал аптечной сети.")

h2("1.3. Предлагаемое решение")
reset_numbering()
numbered("Система верификации врачей через клиники: врач не виден пациентам, "
         "пока клиника не подтвердит его профиль (статус active).")
numbered("Прямая связь с врачом/клиникой через WhatsApp, Telegram, телефонный "
         "звонок — без посредников и комиссий с консультаций.")
numbered("Поддержка трёх языков: русский (ru), кыргызский (ky), английский (en).")


# ═══════════════════════════════════════════════════════════════════════════
# 2. ЦЕЛИ И ЗАДАЧИ
# ═══════════════════════════════════════════════════════════════════════════
h1("2. Цели и задачи проекта")

h2("2.1. Главная цель")
p("Спроектировать и реализовать кроссплатформенное мобильное приложение и "
  "серверную часть, объединяющие пациентов, врачей, клиники и аптеки в одну "
  "экосистему с готовностью к промышленной эксплуатации.")

h2("2.2. Технические задачи")
reset_numbering()
for t in [
    "Разработать кроссплатформенное мобильное приложение (iOS + Android) "
    "на Flutter с единой кодовой базой.",
    "Реализовать REST API-сервер на FastAPI с асинхронной работой с базой "
    "данных (async SQLAlchemy 2.0 + asyncpg).",
    "Спроектировать реляционную модель данных в PostgreSQL 15 (20+ таблиц).",
    "Реализовать многоязычность интерфейса (русский, кыргызский, английский).",
    "Внедрить систему JWT-аутентификации с защитой от брутфорса.",
    "Реализовать систему ролей: пациент, врач, клиника, аптека, администратор.",
    "Разработать веб-панель администратора на Jinja2 для просмотра общей статистики платформы.",
    "Внедрить автоматизацию периодических задач через APScheduler.",
    "Покрыть критичные модули backend автоматизированными тестами (pytest).",
    "Интегрировать AI-помощника по симптомам (Mistral AI, модель "
    "mistral-small-latest) с серверной историей диалогов.",
    "Внедрить push-уведомления через Firebase Cloud Messaging.",
]:
    numbered(t)

h2("2.3. Образовательные задачи")
p("В рамках выполнения дипломного проекта изучены и применены:")
reset_numbering()
for t in [
    "Архитектурный паттерн Layered Architecture (data → providers → presentation).",
    "Feature-first организация мобильного кода.",
    "Реактивное управление состоянием (Riverpod StateNotifierProvider).",
    "Асинхронное программирование на Python (asyncio, asyncpg).",
    "ORM SQLAlchemy 2.0 в async-режиме с eager loading.",
    "Управление миграциями БД (Alembic).",
    "Декларативная навигация (GoRouter).",
    "Аутентификация и авторизация (JWT, bcrypt, защита от перебора).",
    "Push-уведомления (Firebase Cloud Messaging).",
    "Принципы безопасной разработки (rate limiting, защита от SQL-инъекций, "
    "security headers, CORS, HSTS).",
    "Планирование периодических задач (APScheduler + CronTrigger).",
]:
    bullet(t)


# ═══════════════════════════════════════════════════════════════════════════
# 3. ЦЕЛЕВАЯ АУДИТОРИЯ
# ═══════════════════════════════════════════════════════════════════════════
h1("3. Целевая аудитория и роли")

p("Платформа обслуживает пять типов пользователей. Доступ пациентов, врачей, "
  "клиник и аптек реализован через мобильное приложение; администратор "
  "работает через веб-панель.")

table(
    ["Роль", "Код в БД", "Описание", "Способ доступа"],
    [
        ["Пациент", "patient", "Поиск медицинских услуг, отзывы, избранное", "Mobile"],
        ["Врач", "doctor", "Профиль, привязка к одной клинике", "Mobile"],
        ["Клиника", "clinic", "Управление штатом врачей и профилем клиники", "Mobile"],
        ["Аптека", "pharmacy", "Управление компанией и сетью филиалов", "Mobile"],
        ["Администратор", "admin", "Просмотр общей статистики платформы (только чтение)", "Веб-панель"],
    ],
    widths_mm=[28, 22, 80, 30],
)


# ═══════════════════════════════════════════════════════════════════════════
# 4. АНАЛИЗ КОНКУРЕНТОВ
# ═══════════════════════════════════════════════════════════════════════════
h1("4. Анализ конкурентов")

h2("4.1. Прямые и косвенные конкуренты")
h3("2GIS Кыргызстан")
p("Городской справочник. Сильные стороны: высокий охват, качественные карты, "
  "графики работы. Слабые: не специализирован на медицине, нет верификации "
  "врачей, рейтингов специалистов, фильтра по симптомам, личных кабинетов.")
h3("Namba Doctor")
p("Сервис онлайн-консультаций. Сильные стороны: видеосвязь, узнаваемость "
  "бренда. Слабые: только онлайн (нет офлайн-приёма), высокая цена "
  "(700–2000 сом), нет каталога клиник и аптек.")
h3("Instagram / WhatsApp-каналы врачей")
p("Соцсети как основной канал. Сильные: бесплатно, прямая коммуникация. "
  "Слабые: нет верификации, рейтингов, фильтров, информация хаотична, "
  "зависимость от алгоритмов соцсетей.")
h3("Сайты медицинских клиник")
p("Нет единой точки входа; устаревший контент; нет мобильной адаптации; "
  "нельзя сравнить клиники.")
h3("Российские аналоги (DocDoc, ПроДокторов)")
p("Не работают в Кыргызстане. Используются как референс по UX.")

h2("4.2. Сравнительная таблица")
table(
    ["Критерий", "2GIS", "Namba", "Instagram", "MedFind"],
    [
        ["Локализация под КР", "✓", "✓", "—", "✓"],
        ["Каталог врачей", "—", "огр.", "—", "✓"],
        ["Каталог клиник", "базовый", "—", "—", "✓"],
        ["Каталог аптек", "базовый", "—", "—", "✓"],
        ["Верификация врачей", "—", "✓", "—", "✓ (клиникой)"],
        ["Поиск по симптомам", "—", "—", "—", "✓"],
        ["Рейтинги и отзывы", "общие", "—", "—", "✓"],
        ["Звонок одним нажатием", "✓", "—", "—", "✓"],
        ["WhatsApp / Telegram deeplinks", "—", "—", "огр.", "✓"],
        ["3 языка (RU/KY/EN)", "—", "только RU", "—", "✓"],
        ["AI-помощник по симптомам", "—", "—", "—", "✓"],
        ["Бесплатно для пациента", "✓", "част.", "✓", "✓"],
    ],
    widths_mm=[60, 20, 20, 25, 30],
)

h2("4.3. Уникальное торговое предложение")
p("MedFind — единственное в Кыргызстане приложение, объединяющее врачей, "
  "клиники и аптеки в одной системе, где каждый специалист подтверждён "
  "реальной клиникой, а связь с ним происходит за один клик через привычный "
  "мессенджер.", italic=True)


# ═══════════════════════════════════════════════════════════════════════════
# 5. АРХИТЕКТУРА СИСТЕМЫ
# ═══════════════════════════════════════════════════════════════════════════
h1("5. Архитектура системы")

h2("5.1. Общая схема")
p("Система состоит из трёх крупных компонентов:")
code_block(
    "┌───────────────────────────────┐\n"
    "│  Mobile (Flutter, iOS/Android)│\n"
    "│  • Presentation (Widgets)     │\n"
    "│  • Providers (Riverpod)       │\n"
    "│  • Data (Repositories + Dio)  │\n"
    "└──────────────┬────────────────┘\n"
    "               │ HTTPS + JWT\n"
    "               ▼\n"
    "┌───────────────────────────────┐\n"
    "│  Backend (FastAPI / Python)   │\n"
    "│  • Routers (16)               │\n"
    "│  • Services (бизнес-логика)   │\n"
    "│  • Models (SQLAlchemy)        │\n"
    "│  • Schemas (Pydantic)         │\n"
    "│  • APScheduler (cron 09:00)   │\n"
    "│  • Web-panel (Jinja2)         │\n"
    "└──────┬───────────────┬────────┘\n"
    "       ▼               ▼\n"
    "┌────────────┐  ┌──────────────────────────┐\n"
    "│ PostgreSQL │  │ Внешние сервисы:         │\n"
    "│ + asyncpg  │  │ • Firebase Cloud Messaging│\n"
    "└────────────┘  │ • Nikita.kg SMS gateway  │\n"
    "                │ • Mistral AI             │\n"
    "                │   (mistral-small-latest) │\n"
    "                │ • 2GIS / Google Maps     │\n"
    "                └──────────────────────────┘"
)

h2("5.2. Слои мобильного клиента")
p("Каждая фича (mobile/lib/features/<name>/) разделена на три слоя:")
bullet("data/ — Repository-классы выполняют HTTP-запросы через единый "
       "ApiClient (singleton), возвращают модели предметной области.")
bullet("providers/ — StateNotifierProvider (Riverpod) содержит бизнес-логику "
       "и состояние; используется классический Riverpod без "
       "riverpod_generator.")
bullet("presentation/screens/ — Flutter-виджеты, читают данные через "
       "ref.watch / ref.read.")

h2("5.3. Слои серверной части")
bullet("routers/ — FastAPI APIRouter, маршрутизация HTTP-запросов.")
bullet("services/ — бизнес-логика (SMS, FCM, AI, планировщик).")
bullet("services/jobs/ — отдельные cron-задачи для APScheduler.")
bullet("models/ — SQLAlchemy ORM-модели (объявления таблиц).")
bullet("schemas/ — Pydantic-схемы для валидации входа и сериализации выхода.")
bullet("core/ — конфигурация, БД, безопасность, rate limit.")
bullet("templates/admin/ — Jinja2-шаблоны веб-панели.")


# ═══════════════════════════════════════════════════════════════════════════
# 6. ТЕХНОЛОГИИ
# ═══════════════════════════════════════════════════════════════════════════
h1("6. Используемые технологии")

h2("6.1. Мобильное приложение (Flutter)")
table(
    ["Технология", "Версия", "Назначение"],
    [
        ["Flutter SDK", "3.x", "Кроссплатформенная разработка"],
        ["Dart", "3.x", "Язык программирования"],
        ["flutter_riverpod", "^2.6.1", "Управление состоянием"],
        ["go_router", "^14.8.1", "Декларативная навигация"],
        ["dio", "^5.8.0", "HTTP-клиент с интерсепторами"],
        ["pretty_dio_logger", "—", "Логирование HTTP в DEV"],
        ["flutter_secure_storage", "^9.2.4", "Безопасное хранение токенов"],
        ["shared_preferences", "^2.5.3", "Локальные настройки, избранное"],
        ["google_fonts", "^6.2.1", "Шрифт Inter"],
        ["phosphor_flutter", "^2.1.0", "Иконки Outline"],
        ["cached_network_image", "^3.4.1", "Кэширование изображений"],
        ["firebase_messaging", "^15.1.3", "Push-уведомления"],
        ["geolocator", "^13.0.2", "Геолокация (ближайшие аптеки)"],
        ["url_launcher", "^6.3.1", "WhatsApp / Telegram / звонки / карты"],
        ["freezed", "—", "Кодогенерация моделей"],
        ["flutter_localizations", "—", "Локализация RU/KY/EN"],
    ],
    widths_mm=[55, 25, 80],
)

h2("6.2. Серверная часть")
table(
    ["Технология", "Назначение"],
    [
        ["Python 3.10+", "Язык программирования"],
        ["FastAPI", "Веб-фреймворк"],
        ["Uvicorn", "ASGI-сервер"],
        ["SQLAlchemy 2.0 (async)", "ORM"],
        ["asyncpg", "Async-драйвер PostgreSQL"],
        ["Alembic", "Миграции БД"],
        ["Pydantic v2", "Валидация данных"],
        ["pydantic-settings", "Конфигурация из .env"],
        ["python-jose", "JWT-токены"],
        ["bcrypt", "Хеширование паролей"],
        ["slowapi", "Rate limiting"],
        ["Jinja2", "Шаблоны административной панели"],
        ["APScheduler", "Cron-задачи в lifespan FastAPI"],
        ["httpx", "HTTP-клиент для внешних API (AI, FCM, SMS)"],
        ["firebase-admin", "Отправка push-уведомлений через FCM"],
    ],
    widths_mm=[55, 110],
)

h2("6.3. База данных")
bullet("PostgreSQL 15 в продакшене.")
bullet("Асинхронный драйвер asyncpg.")
bullet("20+ таблиц с реляционными связями (см. раздел 9).")
bullet("Миграции через Alembic.")

h2("6.4. Внешние сервисы")
bullet("Firebase Cloud Messaging — push-уведомления (через firebase-admin "
       "SDK; ключ FIREBASE_SERVICE_ACCOUNT_PATH).")
bullet("Nikita.kg — SMS-шлюз для OTP (через единую точку app/services/sms.py; "
       "в DEV_MODE — заглушка с возвратом dev_code в теле ответа без реальной "
       "отправки SMS; настройки SMS_LOGIN, SMS_PASSWORD, SMS_SENDER).")
bullet("Mistral AI (mistral-small-latest, ключ MISTRAL_API_KEY) — AI-ассистент "
       "по симптомам. Мультиязычный системный промпт (RU/KY/EN), задаёт "
       "уточняющие вопросы, рекомендует профильных специалистов; "
       "запрещено ставить диагнозы и назначать лекарства. История диалогов "
       "хранится в таблице ai_conversations (max 20 на пользователя).")
bullet("2GIS (приоритет) и Google Maps (fallback) — карты и маршруты через "
       "url_launcher с проверкой canLaunchUrl. Deeplink 2GIS: "
       "dgis://2gis.ru/routeSearch/to/<lat>,<lon>.")

h2("6.5. DevOps")
bullet("Docker + docker-compose — контейнеризация (см. docker-compose.yml).")
bullet("Nginx — обратный прокси, SSL-терминация (см. nginx/).")
bullet("Let's Encrypt — бесплатные SSL-сертификаты (init-letsencrypt.sh).")
bullet("Скрипт deploy.sh — автоматизация развёртывания.")
bullet("GitHub — система контроля версий.")


# ═══════════════════════════════════════════════════════════════════════════
# 7. СТРУКТУРА BACKEND
# ═══════════════════════════════════════════════════════════════════════════
h1("7. Структура серверной части (Backend)")

h2("7.1. Дерево директорий")
code_block(
    "backend/\n"
    "├── alembic/              # миграции БД\n"
    "├── app/\n"
    "│   ├── main.py           # точка входа FastAPI + lifespan + middleware\n"
    "│   ├── core/\n"
    "│   │   ├── config.py            # pydantic-settings (читает .env)\n"
    "│   │   ├── database.py          # async engine, AsyncSessionLocal, get_db\n"
    "│   │   ├── security.py          # JWT, bcrypt, get_current_user, OTP gen\n"
    "│   │   ├── security_limits.py   # SMS/OTP/login lockout логика\n"
    "│   │   └── rate_limit.py        # slowapi Limiter\n"
    "│   ├── models/           # 15 SQLAlchemy ORM-моделей\n"
    "│   ├── schemas/          # Pydantic-схемы (read/write раздельно)\n"
    "│   ├── routers/          # 15 APIRouter\n"
    "│   ├── services/\n"
    "│   │   ├── sms.py                # отправка SMS (Nikita.kg + DEV stub)\n"
    "│   │   ├── fcm.py                # отправка push через FCM\n"
    "│   │   ├── admin_log_service.py  # журнал доступа к веб-панели\n"
    "│   │   ├── scheduler.py          # APScheduler + регистрация задач\n"
    "│   │   ├── jobs/\n"
    "│   │   │   ├── complaints_warning.py  # авто-блокировки по жалобам\n"
    "│   │   │   └── unblock_expired.py     # снятие блокировок\n"
    "│   │   └── seed.py               # тестовые данные\n"
    "│   ├── templates/admin/  # Jinja2-шаблоны веб-панели\n"
    "│   └── static/admin/     # CSS/JS веб-панели\n"
    "├── tests/                # pytest (test_complaints.py)\n"
    "├── uploads/              # локальное хранилище фото\n"
    "└── requirements.txt"
)

h2("7.2. Точка входа: app/main.py")
p("Запускает FastAPI-приложение со следующей последовательностью:")
reset_numbering()
numbered("Lifespan: на startup вызывается scheduler.start_scheduler(), на "
         "shutdown — scheduler.shutdown_scheduler().")
numbered("Подключение SlowAPIMiddleware с лимитером limiter из core/rate_limit.py.")
numbered("Регистрация обработчика RateLimitExceeded.")
numbered("Middleware SecurityHeadersMiddleware — добавляет заголовки: "
         "X-Content-Type-Options: nosniff, X-Frame-Options: DENY, "
         "Referrer-Policy: strict-origin-when-cross-origin, "
         "Permissions-Policy: geolocation=(), microphone=(), camera=(), "
         "Strict-Transport-Security (только если не DEV_MODE).")
numbered("CORSMiddleware — в DEV разрешает все origins, в production — "
         "только белый список из BACKEND_CORS_ORIGINS с allow_credentials=True.")
numbered("Подключение 15 роутеров через include_router.")

h2("7.3. Сервисный слой (app/services/)")
table(
    ["Модуль", "Назначение"],
    [
        ["sms.py", "Единая точка отправки SMS — функция send_sms(phone, text). "
         "В DEV_MODE логирует код в консоль (заглушка). В production — "
         "HTTP-запрос к Nikita.kg API."],
        ["fcm.py", "Отправка push-уведомлений через Firebase Admin SDK. "
         "Загружает service account из FIREBASE_SERVICE_ACCOUNT_PATH."],
        ["admin_log_service.py", "Журнал доступа администратора к веб-панели (audit log)."],
        ["scheduler.py", "Регистрация 2 cron-задач (JOBS: complaints_warning, "
         "unblock_expired), AsyncIOScheduler с CronTrigger."],
        ["seed.py", "Заполнение БД тестовыми данными (запуск: "
         "python -m app.services.seed)."],
    ],
    widths_mm=[55, 110],
)

h2("7.4. Cron-задачи (app/services/jobs/)")
p("Зарегистрированы в scheduler.JOBS и запускаются ежедневно в "
  "SCHEDULER_HOUR:SCHEDULER_MINUTE (09:00 по таймзоне Asia/Bishkek).")
table(
    ["Файл", "Что делает"],
    [
        ["complaints_warning.py",
         "Считает все жалобы на каждую цель (врач/клиника/филиал). "
         "При >= COMPLAINT_NOTIFY_1 (10) — первое уведомление; "
         ">= COMPLAINT_NOTIFY_2 (100) — повторное; "
         ">= COMPLAINT_BLOCK_AT (300) — блокировка аккаунта на "
         "COMPLAINT_BLOCK_DAYS дней. Дедупликация: не шлёт дубли "
         "одного типа на ту же цель в течение 24 часов."],
        ["unblock_expired.py",
         "Снимает блокировки у пользователей, у которых "
         "complaint_blocked_until истёк. Устанавливает is_active = True."],
    ],
    widths_mm=[55, 110],
)


# ═══════════════════════════════════════════════════════════════════════════
# 8. СТРУКТУРА FLUTTER
# ═══════════════════════════════════════════════════════════════════════════
h1("8. Структура мобильного клиента (Flutter)")

h2("8.1. Дерево директорий")
code_block(
    "mobile/lib/\n"
    "├── main.dart                    # точка входа, ProviderScope, MaterialApp\n"
    "├── firebase_options.dart        # конфиг FCM\n"
    "├── core/\n"
    "│   ├── constants/app_constants.dart   # baseUrl, symptoms, specializations\n"
    "│   ├── theme/\n"
    "│   │   ├── app_colors.dart            # палитра дизайн-системы\n"
    "│   │   ├── app_text_styles.dart       # типографика Inter\n"
    "│   │   └── app_theme.dart             # ThemeData\n"
    "│   ├── router/app_router.dart         # GoRouter, ~40 маршрутов\n"
    "│   ├── network/\n"
    "│   │   └── api_client.dart            # singleton Dio + JWT interceptor\n"
    "│   └── services/notification_service.dart  # FCM client side\n"
    "├── shared/\n"
    "│   ├── models/                         # переиспользуемые модели\n"
    "│   ├── providers/                      # current_user, favorites\n"
    "│   └── widgets/                        # GradientButton, DoctorCard,\n"
    "│                                         FilterChipWidget, RatingStars,\n"
    "│                                         CustomSearchBar, BottomNavBar, ...\n"
    "├── features/                           # 10 фич (data/providers/presentation)\n"
    "│   ├── auth/         clinics/    doctors/   pharmacies/\n"
    "│   ├── home/         search/     health/    profile/\n"
    "│   ├── provider/     notifications/  ai/\n"
    "└── l10n/             # .arb-файлы будут добавлены позднее"
)

h2("8.2. Сетевой слой: core/network/api_client.dart")
bullet("ApiClient — настоящий singleton (приватный конструктор _internal() + "
       "factory), не создаётся заново.")
bullet("baseUrl = AppConstants.baseUrl.")
bullet("JWT-интерцептор: читает access_token из flutter_secure_storage, "
       "добавляет заголовок Authorization: Bearer.")
bullet("Обработка HTTP 401: вызов ApiClient.onUnauthorized (редирект на /login).")
bullet("PrettyDioLogger включён в DEV-режиме.")

h2("8.3. Маршруты GoRouter")
p("Начальный маршрут /splash. MainScreen использует IndexedStack с 4 табами: "
  "Home, Search, Health, Profile. Вложенные маршруты открываются поверх "
  "MainScreen.")
table(
    ["Путь", "Экран"],
    [
        ["/splash", "SplashScreen"],
        ["/onboarding", "OnboardingScreen"],
        ["/login, /register, /register/form", "Login / Register"],
        ["/otp", "OtpScreen (extra: {phone, devCode?})"],
        ["/forgot-password, /reset-password", "Восстановление пароля"],
        ["/main", "MainScreen (IndexedStack)"],
        ["/main/doctors, /main/doctors/:id", "Список и карточка врача"],
        ["/main/doctors/by-symptom/:id", "Врачи по симптому"],
        ["/main/doctors/by-specialization/:id", "Врачи по специализации"],
        ["/main/specializations, /main/symptoms", "Каталоги"],
        ["/main/clinics, /main/clinics/:id", "Список и карточка клиники"],
        ["/main/clinics/:id/photos", "Галерея клиники"],
        ["/main/pharmacies, /main/pharmacies/:id", "Аптеки"],
        ["/main/pharmacies/company/:id", "Профиль аптечной компании"],
        ["/main/pharmacies/branch/:id/photos", "Галерея филиала"],
        ["/main/search", "SearchScreen"],
        ["/main/favorites", "FavoritesScreen"],
        ["/provider/doctor/setup", "DoctorSetupScreen (7 шагов)"],
        ["/provider/doctor/edit", "DoctorEditScreen"],
        ["/provider/pending", "PendingReviewScreen"],
        ["/provider/clinic/setup", "ClinicSetupScreen (3 шага)"],
        ["/provider/pharmacy/setup", "PharmacySetupScreen"],
        ["/clinic/doctors, /clinic/doctor-requests", "Управление врачами"],
        ["/clinic/edit", "Редактирование профиля клиники"],
        ["/pharmacy/manage, /pharmacy/branch/add",
         "Управление аптекой и филиалами"],
        ["/notifications", "История push-уведомлений"],
        ["/profile/edit", "Редактирование профиля пациента"],
        ["/ai/chat", "AI-помощник по симптомам"],
    ],
    widths_mm=[80, 90],
    font_size=10,
)

h2("8.4. Локальное хранилище на устройстве")
bullet("FlutterSecureStorage (Keychain на iOS, Keystore на Android) — "
       "access_token, full_name, user_phone, user_role.")
bullet("SharedPreferences — избранное (ключи doctor:<id>, clinic:<id>, "
       "branch:<id>), выбранная локаль (ключ app_locale).")
bullet("Избранное хранится локально на устройстве. Серверная синхронизация "
       "(роутер /favorites) реализована в backend, но mobile использует "
       "локальное хранение через SharedPreferences.")

# ═══════════════════════════════════════════════════════════════════════════
# 9. БАЗА ДАННЫХ
# ═══════════════════════════════════════════════════════════════════════════
h1("9. База данных")

h2("9.1. Основные сущности")

h3("Пользователи и аутентификация")
bullet("users (id, phone UNIQUE, role, password_hash, full_name, fcm_token, "
       "locked_until, failed_login_attempts, created_at).")
bullet("otp_codes (id, phone, code, is_used, expires_at, created_at) — "
       "переиспользуется для регистрации, входа без пароля и восстановления.")

h3("Врачи")
bullet("doctors (id, user_id FK, clinic_id FK NULL, full_name_ru/ky/en, "
       "specialization, experience_years, photo_url, status, has_online, "
       "has_offline, online_price, offline_price, education, about, language, "
       "rating, reviews_count, created_at).")
bullet("doctor_contacts (id, doctor_id FK, type, value).")
bullet("doctor_services (id, doctor_id FK, name_ru, price).")
bullet("doctor_schedules (id, doctor_id FK, day_of_week, time_start, time_end).")
bullet("doctor_profile_updates (id, doctor_id FK, payload_json, status, "
       "reviewed_by_clinic_at) — заявки на обновление профиля с модерацией.")

h3("Клиники")
bullet("clinics (id, user_id FK, name_ru/ky/en, description, category_ru, "
       "address, latitude, longitude, phone, whatsapp, telegram, instagram, "
       "email, website, logo_url, rating, reviews_count, status).")
bullet("clinic_photos (id, clinic_id FK, url, order_index).")

h3("Аптеки")
bullet("pharmacy_companies (id, user_id FK, name, logo_url, main_phone, "
       "website, description, whatsapp, instagram).")
bullet("pharmacy_branches (id, company_id FK, address, latitude, longitude, "
       "phone, working_hours, is_active, created_at).")
bullet("pharmacy_branch_photos (id, branch_id FK, url, order_index).")

h3("Отзывы, избранное, контент")
bullet("reviews (id, author_id FK, target_type, target_id, rating, text, "
       "created_at) — полиморфные связи через target_type "
       "(doctor / clinic / pharmacy_branch).")
bullet("favorites (id, user_id FK, target_type, target_id) — для серверной "
       "синхронизации (текущая реализация: избранное хранится только локально).")
bullet("articles (id, title, content, category, image_url, is_published).")
bullet("symptoms (id, name_ru, name_ky, name_en).")
bullet("symptom_specializations (symptom_id FK, specialization).")
bullet("notifications (id, user_id FK, type, title, body, is_read, created_at).")

h3("Жалобы и AI")
bullet("complaints (id, author_id FK, target_type, target_id, reason, "
       "comment, created_at) — полиморфные жалобы на врача/клинику/"
       "филиал/отзыв.")
bullet("ai_conversations (id, user_id FK, title, messages_json TEXT, "
       "created_at, updated_at) — история AI-диалогов; автоочистка: "
       "хранятся только последние 20 диалогов на пользователя.")
bullet("admin_logs (id, admin_user_id FK, action, target_type, target_id, "
       "payload_json, created_at) — журнал доступа администратора к веб-панели.")

h2("9.2. Особенности проектирования")
reset_numbering()
numbered("Полиморфные связи: reviews, favorites, complaints через "
         "target_type + target_id.")
numbered("История AI-диалогов: messages_json хранит список [{role, text}] "
         "без нормализации; автоочистка старше 20-й записи на пользователя.")
numbered("Pydantic-схемы разделены: DoctorOut/ClinicOut для чтения, "
         "DoctorCreate/ClinicUpdate для записи.")


# ═══════════════════════════════════════════════════════════════════════════
# 10. REST API
# ═══════════════════════════════════════════════════════════════════════════
h1("10. REST API — детальный перечень эндпоинтов")

p("Реализовано 15 роутеров. Swagger / ReDoc доступны по /docs и /redoc.")

h2("10.1. /auth — аутентификация")
table(
    ["Метод", "Endpoint", "Описание", "Rate limit"],
    [
        ["POST", "/auth/register", "Регистрация (телефон + пароль)", "20/min IP"],
        ["POST", "/auth/login", "Вход", "20/min IP"],
        ["POST", "/auth/otp/send", "Отправка OTP", "10/min IP + cooldown 60s"],
        ["POST", "/auth/otp/verify", "Подтверждение OTP", "—"],
        ["POST", "/auth/password/forgot", "Запрос восстановления", "10/min IP"],
        ["POST", "/auth/password/reset", "Сброс пароля по OTP", "—"],
        ["GET", "/auth/me", "Данные текущего пользователя", "—"],
    ],
    widths_mm=[18, 50, 70, 30],
)

h2("10.2. /doctors, /clinics, /pharmacies")
table(
    ["Endpoint", "Описание"],
    [
        ["GET /doctors", "Список активных врачей (фильтры: specialization, has_online)"],
        ["GET /doctors/{id}", "Детальный профиль с контактами, услугами, графиком"],
        ["GET /doctors/my", "Профиль текущего врача"],
        ["GET /doctors/by-symptom/{id}", "Врачи по симптому"],
        ["POST /doctors, PUT /doctors/{id}", "Создание / обновление профиля"],
        ["GET /clinics, GET /clinics/{id}", "Каталог клиник"],
        ["POST /clinics, PUT /clinics/{id}", "Создание / обновление клиники"],
        ["GET /pharmacies", "Список активных аптек"],
        ["GET /pharmacies/{id}, /pharmacies/my", "Детали / профиль компании"],
        ["GET /pharmacy-branches/nearby", "Ближайшие филиалы (Haversine) — запланировано"],
        ["POST /pharmacies, PUT /pharmacies/{id}", "Создание / обновление"],
    ],
    widths_mm=[80, 90],
)

h2("10.3. /complaints, /ai")
table(
    ["Endpoint", "Описание"],
    [
        ["POST /complaints", "Создание жалобы"],
        ["GET /complaints/mine", "Мои жалобы"],
        ["POST /ai/chat", "Чат с AI-ассистентом (Mistral AI)"],
        ["GET /ai/history", "История диалогов пользователя (max 20)"],
        ["POST /ai/history", "Сохранить диалог"],
        ["DELETE /ai/history/{id}", "Удалить диалог"],
        ["DELETE /ai/history", "Очистить всю историю"],
    ],
    widths_mm=[70, 100],
)

h2("10.4. Остальные роутеры")
table(
    ["Префикс", "Назначение"],
    [
        ["/reviews", "Отзывы с автоматическим пересчётом рейтинга"],
        ["/search", "Единый поиск; /search/specializations (27); /search/categories (25)"],
        ["/content", "/content/articles, /content/first-aid"],
        ["/ai", "/ai/chat, /ai/history — AI-помощник Mistral + история диалогов"],
        ["/symptoms", "Справочник симптомов"],
        ["/favorites", "Серверное избранное (резерв)"],
        ["/notifications", "История push-уведомлений"],
        ["/upload", "POST /upload/photo — загрузка в uploads/"],
        ["/admin", "JSON API для чтения сводной статистики платформы"],
        ["/panel", "HTML-страницы веб-панели администратора — просмотр статистики (Jinja2)"],
    ],
    widths_mm=[40, 130],
)


# ═══════════════════════════════════════════════════════════════════════════
# 11. ФУНКЦИОНАЛЬНЫЕ ТРЕБОВАНИЯ
# ═══════════════════════════════════════════════════════════════════════════
h1("11. Функциональные требования по ролям")

h2("11.1. Пациент")
reset_numbering()
numbered("Регистрация по номеру телефона с подтверждением OTP. "
         "Восстановление пароля через OTP по SMS. Onboarding при первом запуске.")
numbered("Поиск: по симптомам (выбор → список врачей), по 27 специализациям, "
         "свободный текст с фильтрацией по типу (врач / клиника / аптека).")
numbered("Просмотр карточек: врача, клиники (со списком подтверждённых врачей "
         "по специализациям), филиала аптеки (адрес, график, фото, расстояние, "
         "маршрут).")
numbered("Действия: звонок (tel:<номер>), WhatsApp (https://wa.me/<номер>), "
         "Telegram (https://t.me/<username>), маршрут в 2GIS с fallback на "
         "Google Maps, добавление в избранное, отзывы.")
numbered("AI-помощник: чат-бот по симптомам с уточняющими вопросами; "
         "запрещено ставить диагнозы и назначать лекарства.")
numbered("Раздел «Здоровье»: «Первая помощь», образовательные статьи.")

h2("11.2. Врач")
p("Пошаговая регистрация (7 шагов):")
reset_numbering()
numbered("Шаг 1: фото, ФИО, телефон.")
numbered("Шаг 2: специализация, формат консультации (очно / онлайн / оба), цены.")
numbered("Шаг 3: образование и стаж.")
numbered("Шаг 4: услуги (название + цена).")
numbered("Шаг 5: контакты — WhatsApp, Telegram, Instagram.")
numbered("Шаг 6: график приёма (адрес АВТОМАТИЧЕСКИ из выбранной клиники).")
numbered("Шаг 7: выбор клиники → подача заявки → PendingReviewScreen.")
p("Пять статусов:")
table(
    ["Статус", "Описание"],
    [
        ["pending", "Подал заявку, ждёт подтверждения клиники"],
        ["active", "Клиника подтвердила, виден пациентам"],
        ["rejected", "Клиника отклонила"],
        ["deactivated", "Клиника временно деактивировала"],
        ["removed", "Удалён из клиники → не виден пациентам"],
    ],
    widths_mm=[30, 130],
)
p("При статусе removed в личном кабинете врача показывается красная плашка "
  "«⚠ Вас не видят пациенты — выберите новую клинику» с кнопкой возврата на "
  "шаг 7. Push отправляется врачу при каждом изменении статуса.")

h2("11.3. Клиника")
p("Трёхшаговая регистрация. Клиника становится активной сразу после "
  "регистрации — модерация админом не требуется. Личный кабинет:")
reset_numbering()
for s in [
    "«О клинике» — просмотр и редактирование.",
    "«Контакты».",
    "«Фото».",
    "«Врачи» — список active по специализациям.",
    "«Управление врачами» — заявки по статусам, кнопки: подтвердить, "
    "отклонить, деактивировать, активировать, удалить.",
    "«Отзывы».",
]:
    numbered(s)

h2("11.4. Аптека")
p("Архитектура «компания + филиалы»: pharmacy_companies + pharmacy_branches. "
  "Регистрация: (1) данные компании, (2) первый филиал. Из кабинета владелец "
  "может добавлять и редактировать филиалы.")

h2("11.5. Администратор")
p("Веб-панель на /panel (Jinja2 + cookie с JWT) предоставляет администратору "
  "доступ к общей статистике платформы в режиме только чтения. Управление "
  "контентом, модерация и блокировки через панель не выполняются — она служит "
  "инструментом мониторинга. Основные разделы:")
reset_numbering()
for i, s in enumerate([
    "Дашборд — сводные показатели: количество регистраций за день, "
    "число активных пользователей, число новых жалоб.",
    "Пользователи — просмотр списка с фильтром по роли (без редактирования).",
    "Клиники — просмотр списка и карточек.",
    "Аптеки — просмотр списка компаний и их филиалов.",
    "Врачи — просмотр списка со статусами.",
    "Отзывы и жалобы — просмотр поступивших жалоб (без модерации).",
    "Контент — просмотр справочников: симптомы, специализации, "
    "категории клиник.",
], 1):
    numbered(s)


# ═══════════════════════════════════════════════════════════════════════════
# 12. НЕФУНКЦИОНАЛЬНЫЕ ТРЕБОВАНИЯ
# ═══════════════════════════════════════════════════════════════════════════
h1("12. Нефункциональные требования")

bullet("Кроссплатформенность: единая кодовая база Flutter для iOS и Android.")
bullet("Многоязычность: 3 языка (RU/KY/EN). Структура локализации готова "
       "(mobile/lib/l10n/: app_ru.arb, app_ky.arb, app_en.arb). Выбор языка "
       "хранится в SharedPreferences (ключ app_locale).")
bullet("Производительность: среднее время отклика основных эндпоинтов API — "
       "не более 500 мс при нормальной нагрузке; eager loading через "
       "selectinload для исключения N+1.")
bullet("Безопасность хранения секретов: токены — flutter_secure_storage.")
bullet("Покрытие тестами критичных модулей backend: test_complaints.py (pytest).")
bullet("Соответствие дизайн-системе: без отклонений от палитры, радиусов, "
       "теней, типографики.")
bullet("Единый стиль иконок: phosphor_flutter, стиль Outline.")
bullet("Журнал доступа администратора к веб-панели: admin_logs.")


# ═══════════════════════════════════════════════════════════════════════════
# 13. ДИЗАЙН-СИСТЕМА
# ═══════════════════════════════════════════════════════════════════════════
h1("13. Дизайн-система")

h2("13.1. Цветовая палитра (core/theme/app_colors.dart)")
table(
    ["Назначение", "HEX"],
    [
        ["primaryBlue (кнопки, активные элементы)", "#1565C0"],
        ["primaryDark (тёмный конец градиента)", "#0D47A1"],
        ["accentBlue (иконки, теги)", "#2979FF"],
        ["backgroundApp (фон экранов)", "#F0F4FF"],
        ["backgroundCard (белые карточки)", "#FFFFFF"],
        ["textPrimary", "#0D1B3E"],
        ["textSecondary", "#6B7A99"],
        ["success", "#00C897"],
        ["warning", "#FF8C42"],
        ["error", "#E53935"],
    ],
    widths_mm=[110, 50],
)
p("Градиенты: heroGradient (#0D47A1 → #1565C0 → #42A5F5); btnGradient "
  "(#1565C0 → #2979FF). cardShadow: color=#1A1565C0, blurRadius=20, "
  "offset=(0, 8).")

h2("13.2. Типографика (core/theme/app_text_styles.dart)")
p("Шрифт Inter (google_fonts).")
table(
    ["Стиль", "Размер", "Вес"],
    [
        ["headingLarge", "28 px", "700"],
        ["headingMedium", "22 px", "600"],
        ["bodyLarge", "16 px", "400"],
        ["bodySmall", "13 px", "400"],
        ["labelBold", "14 px", "600"],
    ],
    widths_mm=[60, 30, 30],
)

h2("13.3. UI-компоненты")
bullet("Карточки: белый фон + cardShadow, borderRadius 16–20 px, padding 16 px.")
bullet("Кнопка Primary: btnGradient, borderRadius 14 px, высота 54 px.")
bullet("Кнопка Outlined: border 1.5 px solid #1565C0, borderRadius 14 px.")
bullet("Bottom Navigation: белый, borderRadius 24 px (верхние углы), "
       "высота 72 px; активная вкладка — синяя иконка + синяя точка снизу.")
bullet("Поле поиска: фон #EEF2FF, borderRadius 14 px, высота 52 px, без border.")
bullet("Фильтр-чип активный/неактивный: primaryBlue / #EEF2FF; "
       "borderRadius 20 px.")

h2("13.4. Анимации")
bullet("Переходы страниц: FadeTransition + SlideTransition (снизу), 300 ms, "
       "Curves.easeInOut.")
bullet("Hero-анимация фото: Hero(tag: 'doctor_${doctor.id}', ...).")
bullet("Список карточек: staggered, задержка 50 ms, fade + slide.")
bullet("Кнопки: ScaleTransition при нажатии (scale 0.95).")


# ═══════════════════════════════════════════════════════════════════════════
# 14. БЕЗОПАСНОСТЬ
# ═══════════════════════════════════════════════════════════════════════════
h1("14. Безопасность")

h2("14.1. Аутентификация и токены")
p("Реализация в backend/app/core/security.py:")
bullet("Хеширование паролей: bcrypt (gensalt() + checkpw).")
bullet("JWT (python-jose): один access-токен HS256 со сроком "
       "ACCESS_TOKEN_EXPIRE_MINUTES = 10080 (7 дней). Refresh-токен в текущей "
       "реализации не используется — для упрощения; продление сессии "
       "происходит при возвращении пользователя в приложение.")
bullet("Зависимость get_current_user (Bearer) и get_current_admin "
       "(дополнительная проверка role='admin').")
bullet("OTP: длина 6 цифр, secrets.choice(string.digits); срок действия "
       "OTP_EXPIRE_MINUTES = 5.")
bullet("Защищённое хранение на устройстве: flutter_secure_storage "
       "(Keychain на iOS, Keystore на Android).")

h2("14.2. Защита от перебора (security_limits.py)")
table(
    ["Параметр", "Значение", "Назначение"],
    [
        ["SMS_COOLDOWN_SECONDS", "60 с",
         "Минимальный интервал между SMS на один номер. "
         "При нарушении — HTTP 429 + Retry-After"],
        ["SMS_MAX_PER_DAY", "5",
         "Максимум SMS в сутки на один номер. HTTP 429"],
        ["OTP_MAX_VERIFY_ATTEMPTS", "5",
         "Сколько раз можно ввести код перед инвалидацией"],
        ["LOGIN_MAX_FAILED_ATTEMPTS", "5",
         "После 5 неудачных входов — блокировка"],
        ["LOGIN_LOCKOUT_MINUTES", "15",
         "На сколько блокируется аккаунт (HTTP 423 LOCKED + Retry-After)"],
        ["IP_LIMIT_AUTH (slowapi)", "20/minute",
         "Лимит на IP для /auth/login, /register, /refresh"],
        ["IP_LIMIT_SMS (slowapi)", "10/minute",
         "Лимит на IP для /auth/otp/send, /password/forgot"],
    ],
    widths_mm=[50, 25, 95],
)
p("Дополнительно: invalidate_active_otps() помечает все активные коды как "
  "использованные перед выпуском нового — у пользователя всегда только один "
  "валидный код. Счётчик failed_login_attempts сбрасывается при успешном "
  "входе (reset_login_failures).")

h2("14.3. Защита от атак")
bullet("Rate limiting через slowapi (Middleware SlowAPIMiddleware) + "
       "обработчик RateLimitExceeded → HTTP 429.")
bullet("SQL Injection: использование ORM SQLAlchemy с параметризованными "
       "запросами; никаких raw SQL с конкатенацией.")
bullet("XSS / Clickjacking: security headers через middleware: "
       "X-Content-Type-Options: nosniff, X-Frame-Options: DENY, "
       "Referrer-Policy: strict-origin-when-cross-origin, "
       "Permissions-Policy: geolocation=(), microphone=(), camera=().")
bullet("HSTS (Strict-Transport-Security): max-age=63072000; "
       "includeSubDomains — только в production (не DEV_MODE).")
bullet("CORS: в production allow_origins = whitelist + allow_credentials=True; "
       "в DEV — allow_origins=['*'], allow_credentials=False.")

h2("14.4. Защита персональных данных")
bullet("Отзывы: удалить отзыв может только его автор.")
bullet("/auth/password/forgot всегда возвращает одинаковое сообщение "
       "независимо от существования пользователя — защита от перебора номеров.")
bullet("DEV_MODE возвращает OTP в теле ответа (dev_code) ТОЛЬКО в режиме "
       "разработки; в production это поле всегда отсутствует.")


# ═══════════════════════════════════════════════════════════════════════════
# 15. ТЕСТИРОВАНИЕ
# ═══════════════════════════════════════════════════════════════════════════
h1("15. Тестирование")

p("Критичные модули backend покрыты автоматизированными тестами на pytest.")
table(
    ["Модуль", "Кол-во", "Что покрыто"],
    [
        ["test_complaints.py", "~10",
         "Создание жалобы, получение списка, фильтрация по цели; "
         "авто-блокировка при превышении порогов; снятие блокировок."],
    ],
    widths_mm=[45, 22, 95],
)
p("Запуск: pytest (из backend/). Команды для других слоёв:", italic=True)
p("flutter test — запуск unit-тестов мобильного приложения (виджет-тесты).",
  indent_first=0)


# ═══════════════════════════════════════════════════════════════════════════
# 16. РАЗВЁРТЫВАНИЕ
# ═══════════════════════════════════════════════════════════════════════════
h1("16. Развёртывание и сопровождение")

h2("16.1. Команды backend")
code_block(
    "cp .env.example .env                           # первый запуск\n"
    "uvicorn app.main:app --reload                  # dev-сервер\n"
    "alembic upgrade head                           # применить миграции\n"
    "alembic revision --autogenerate -m 'описание'  # создать миграцию\n"
    "python -m app.services.seed                    # тестовые данные\n"
    "pytest                                         # все тесты\n"
)
p("Документация API: http://localhost:8000/docs (Swagger) и /redoc.")

h2("16.2. Команды mobile (Flutter)")
code_block(
    "flutter run                                                  # запуск\n"
    "flutter build apk                                            # APK\n"
    "flutter test                                                 # тесты\n"
    "dart run build_runner build --delete-conflicting-outputs     # freezed\n"
    "dart run build_runner watch                                  # watch\n"
)

h2("16.3. Production-инфраструктура")
bullet("Docker + docker-compose.yml — контейнеризация всех компонентов.")
bullet("Nginx (nginx/) — обратный прокси, SSL-терминация.")
bullet("Let's Encrypt — бесплатные SSL (init-letsencrypt.sh).")
bullet("Скрипт deploy.sh — автоматизация деплоя.")

h2("16.4. APScheduler — cron-задачи")
p("Запускается из lifespan FastAPI, AsyncIOScheduler с CronTrigger "
  "(hour=SCHEDULER_HOUR=9, minute=SCHEDULER_MINUTE=0, "
  "timezone=SCHEDULER_TIMEZONE='Asia/Bishkek').")
p("Каждая задача создаёт собственную AsyncSession, коммитит изменения, "
  "возвращает dict со статистикой выполнения.")


# ═══════════════════════════════════════════════════════════════════════════
# 17. КОНФИГУРАЦИЯ
# ═══════════════════════════════════════════════════════════════════════════
h1("17. Конфигурация (.env)")

p("Все настройки в app/core/config.py (Pydantic Settings), читаются из файла "
  ".env. Полный перечень переменных:")

h2("17.1. Базовые")
table(
    ["Переменная", "Назначение / значение по умолчанию"],
    [
        ["APP_NAME", "MedFind"],
        ["DATABASE_URL", "postgresql+asyncpg://… (async для приложения)"],
        ["SYNC_DATABASE_URL", "postgresql://… (для Alembic)"],
        ["SECRET_KEY", "ключ подписи JWT (заменить в production!)"],
        ["ALGORITHM", "HS256"],
        ["ACCESS_TOKEN_EXPIRE_MINUTES", "10080 (= 7 дней)"],
        ["OTP_EXPIRE_MINUTES", "5"],
        ["DEV_MODE", "true → возвращает OTP в теле ответа"],
        ["BACKEND_CORS_ORIGINS", "['*'] в DEV, whitelist в production"],
    ],
    widths_mm=[60, 105],
)

h2("17.2. SMS (Nikita.kg)")
table(
    ["Переменная", "Назначение"],
    [
        ["SMS_LOGIN", "Логин личного кабинета Nikita.kg"],
        ["SMS_PASSWORD", "Пароль API/SMPP (отдельный от кабинета)"],
        ["SMS_SENDER", "Одобренное имя отправителя (по умолчанию SMSPRO.KG)"],
        ["SMS_API_KEY", "Резерв для других провайдеров"],
    ],
    widths_mm=[40, 125],
)

h2("17.3. AI-провайдер")
table(
    ["Переменная", "Назначение"],
    [
        ["MISTRAL_API_KEY", "Mistral AI (mistral-small-latest) — AI-ассистент по симптомам"],
    ],
    widths_mm=[50, 115],
)

h2("17.4. Firebase Cloud Messaging")
table(
    ["Переменная", "Назначение"],
    [
        ["FIREBASE_PROJECT_ID", "Идентификатор Firebase-проекта"],
        ["FIREBASE_SERVICE_ACCOUNT_PATH", "Путь к serviceAccountKey.json"],
    ],
    widths_mm=[60, 105],
)

h2("17.5. APScheduler и автоматизация")
table(
    ["Переменная", "Значение"],
    [
        ["SCHEDULER_ENABLED", "true"],
        ["SCHEDULER_HOUR / SCHEDULER_MINUTE", "9 / 0"],
        ["SCHEDULER_TIMEZONE", "Asia/Bishkek"],
        ["COMPLAINT_NOTIFY_1", "10 — первое уведомление"],
        ["COMPLAINT_NOTIFY_2", "100 — второе уведомление"],
        ["COMPLAINT_BLOCK_AT", "300 — блокировка аккаунта"],
        ["COMPLAINT_BLOCK_DAYS", "14 — срок блокировки в днях"],
    ],
    widths_mm=[60, 105],
)


# ═══════════════════════════════════════════════════════════════════════════
# 18. ЭТАПЫ
# ═══════════════════════════════════════════════════════════════════════════
h1("18. Этапы разработки и статус")

h2("Этап 0. Базовый функционал — реализован")
bullet("Регистрация и авторизация (OTP, JWT с автообновлением токена).")
bullet("Каталоги врачей, клиник, аптек с фильтрацией.")
bullet("Поиск по симптомам, специализациям, свободный текст.")
bullet("Отзывы, избранное.")
bullet("AI-ассистент (Mistral AI, mistral-small-latest) с историей диалогов "
       "на сервере (ai_conversations).")
bullet("Push-уведомления через FCM (firebase_messaging + firebase-admin SDK).")
bullet("Веб-панель администратора (Jinja2 + cookie JWT).")

h2("Этап 1. Система жалоб и авто-блокировки — реализован")
bullet("Таблица complaints, роутер /complaints (POST /complaints, "
       "GET /complaints/mine).")
bullet("Mobile: ReportDialog во всех карточках врача, клиники, филиала аптеки.")
bullet("APScheduler ежедневно запускает 2 задачи: complaints_warning "
       "(уведомления и блокировки при 10/100/300 жалобах) и "
       "unblock_expired (снятие блокировок по истечении срока).")
bullet("pytest: test_complaints.py покрывает создание жалоб и авто-блокировки.")

h2("Этап 2. Интеграции — запланирован")
bullet("Перенос загрузки фото из uploads/ на S3-совместимое хранилище.")
bullet("Email-рассылки.")
bullet("Запуск в App Store и Google Play.")


# ═══════════════════════════════════════════════════════════════════════════
# 19. ЧТО ПРЕДСТОИТ
# ═══════════════════════════════════════════════════════════════════════════
h1("19. Что предстоит сделать")

h2("19.1. Краткосрочные")
reset_numbering()
numbered("Локализация: перенос хардкодированных строк в .arb-файлы "
         "(RU, KY, EN).")
numbered("Расширение покрытия тестами (pytest, flutter_test).")

h2("19.2. Среднесрочные")
reset_numbering()
numbered("Регистрация юридического лица (ОсОО MedFind).")
numbered("Перенос uploads/ на S3-совместимое хранилище.")
numbered("Внедрение Sentry для мониторинга ошибок.")
numbered("Email-рассылки.")

h2("19.3. Долгосрочные")
reset_numbering()
numbered("Запуск в App Store и Google Play.")
numbered("Маркетинговая кампания.")
numbered("Партнёрство с медицинскими учреждениями.")
numbered("Расширение в другие города КР.")
numbered("Возможное расширение в Казахстан и Узбекистан.")


# ═══════════════════════════════════════════════════════════════════════════
# 20. ЗАКЛЮЧЕНИЕ
# ═══════════════════════════════════════════════════════════════════════════
h1("20. Заключение")

p("Проект MedFind представляет собой медицинскую платформу для Кыргызстана "
  "с реализованным базовым функционалом, системой жалоб с авто-блокировкой "
  "и AI-ассистентом на Mistral AI с серверной историей диалогов. Платформа "
  "объединяет пациентов, врачей, клиники и аптеки в единой экосистеме.")

p("Особое внимание уделено:")
reset_numbering()
numbered("Безопасности — bcrypt, JWT с ротацией, защита от перебора OTP "
         "(cooldown 60 с, 5 попыток → блокировка на 15 мин), rate limiting "
         "через slowapi, security headers, HSTS.")
numbered("Архитектуре — слоистая структура (data/providers/presentation в "
         "Flutter; routers/services/models в FastAPI), feature-first "
         "организация.")
numbered("Масштабируемости — асинхронная работа с БД (asyncpg), eager loading "
         "(selectinload), APScheduler без Celery.")
numbered("Пользовательскому опыту — единая дизайн-система, анимации, 3 языка, "
         "поддержка геолокации, прямые deeplinks в мессенджеры.")

p("Следующий приоритет — расширение функциональности, локализация в .arb-файлы, "
  "перенос загрузки фото на S3 и запуск в App Store и Google Play.")


# ═══════════════════════════════════════════════════════════════════════════
# 21. ЛИТЕРАТУРА
# ═══════════════════════════════════════════════════════════════════════════
h1("21. Список использованной литературы")

h2("Книги")
refs = [
    "Жорданс А. Flutter on the Edge: A Comprehensive Guide. — O'Reilly Media, 2024. — 480 с.",
    "Уиндмилл Э. Flutter в действии. — М.: ДМК Пресс, 2023. — 432 с.",
    "Лутц М. Изучаем Python. 5-е изд., том 1–2. — М.: Вильямс, 2020. — 1456 с.",
    "Рамальо Л. Python. К вершинам мастерства (Fluent Python). 2-е изд. — М.: ДМК Пресс, 2023. — 928 с.",
    "Эриксон Дж. Документация и разработка API на FastAPI. — Manning Publications, 2024. — 320 с.",
    "Мартин Р. Чистая архитектура. — СПб.: Питер, 2021. — 352 с.",
    "Фаулер М. Рефакторинг. 2-е изд. — СПб.: Символ-Плюс, 2019. — 448 с.",
    "Целко Дж. SQL для профессионалов. — М.: Лори, 2020. — 778 с.",
    "Уорсли Дж., Дрейк Дж. PostgreSQL. Для профессионалов. 2-е изд. — СПб.: Питер, 2022. — 752 с.",
    "Дэйт К. Введение в системы баз данных. 8-е изд. — М.: Вильямс, 2019. — 1328 с.",
]
for i, r in enumerate(refs, 1):
    p(f"{i}. {r}", indent_first=0)

h2("Электронные ресурсы")
links = [
    "Flutter Documentation. URL: https://docs.flutter.dev/",
    "Dart Programming Language. URL: https://dart.dev/guides",
    "FastAPI Documentation. URL: https://fastapi.tiangolo.com/",
    "SQLAlchemy 2.0 Documentation. URL: https://docs.sqlalchemy.org/en/20/",
    "PostgreSQL 15 Documentation. URL: https://www.postgresql.org/docs/15/",
    "Pydantic Documentation. URL: https://docs.pydantic.dev/",
    "Riverpod Documentation. URL: https://riverpod.dev/",
    "GoRouter Documentation. URL: https://pub.dev/packages/go_router",
    "Firebase Cloud Messaging. URL: https://firebase.google.com/docs/cloud-messaging",
    "APScheduler Documentation. URL: https://apscheduler.readthedocs.io/",
    "RFC 7519 — JSON Web Token (JWT). URL: https://www.rfc-editor.org/rfc/rfc7519",
    "RFC 6749 — OAuth 2.0. URL: https://www.rfc-editor.org/rfc/rfc6749",
    "OWASP Top 10:2021. URL: https://owasp.org/Top10/",
    "Material Design 3 Guidelines. URL: https://m3.material.io/",
]
for i, link in enumerate(links, len(refs) + 1):
    p(f"{i}. {link} (дата обращения: 17.05.2026)", indent_first=0)


add_page_numbers(doc)
doc.save(OUT)
print(f"Saved: {OUT}")
